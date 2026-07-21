#!/usr/bin/env python3
"""生成 PsyClaw 使用白皮书的 Word 版(风格对齐用户提供的参考白皮书)。

参考件结构:封面(主标题/副标题/破折号行/定位行/日期版次)→ 关于本白皮书 →
如何使用本白皮书 → 执行摘要 → 目录(域代码,Word 打开后 F9 更新)→ 第 N 章 → 附录。
信息框三色:蓝=命令/提示词示例、绿=实用建议、黄=风险警示;灰底=代码。

内容真源是 docs/WHITEPAPER.md 所依据的实测数据;本脚本只负责排版。
用法:uv run --python 3.12 --with python-docx python scripts/build_whitepaper_docx.py
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
import psyclaw  # noqa: E402

VER = psyclaw.__version__
CN_FONT = "Songti SC"          # 正文中文
CN_HEI = "Heiti SC"            # 标题中文
MONO = "Menlo"

BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x1E, 0x6B, 0x3A)
AMBER = RGBColor(0x8A, 0x5A, 0x00)
GREY = RGBColor(0x44, 0x44, 0x44)


# ── 底层排版助手 ────────────────────────────────────────────────────────────
def _cn(run, font=CN_FONT):
    """中文字体必须同时设 eastAsia,否则 Word 里中文回落成宋体默认、字重不对。"""
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def _shade(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def para(doc, text="", size=10.5, bold=False, color=None, align=None,
         space_after=6, font=CN_FONT, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    _cn(r, font)
    return p


def heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    h.paragraph_format.space_after = Pt(8)
    r = h.add_run(text)
    r.font.size = Pt(18 if level == 1 else 14 if level == 2 else 12)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    r.bold = True
    _cn(r, CN_HEI)
    return h


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(text)
    r.font.size = Pt(size)
    _cn(r)
    return p


def numbered(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(text)
    r.font.size = Pt(size)
    _cn(r)
    return p


def box(doc, title, lines, kind="blue"):
    """三色信息框:blue=命令/示例 green=实用建议 amber=风险警示。"""
    fill, color = {"blue": ("EAF1F8", BLUE), "green": ("EAF5EE", GREEN),
                   "amber": ("FBF3E2", AMBER)}[kind]
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    _shade(cell, fill)
    cell.paragraphs[0].text = ""
    ph = cell.paragraphs[0]
    ph.paragraph_format.space_after = Pt(2)
    hr = ph.add_run(title)
    hr.bold = True
    hr.font.size = Pt(10)
    hr.font.color.rgb = color
    _cn(hr, CN_HEI)
    for ln in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(ln)
        r.font.size = Pt(9.5)
        _cn(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def code(doc, lines):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _shade(cell, "F2F2F2")
    cell.paragraphs[0].text = ""
    first = True
    for ln in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(ln)
        r.font.name = MONO
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        _shade(c, "E8EDF3")
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        _cn(r, CN_HEI)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].paragraphs[0].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9)
            _cn(r)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def toc_field(doc):
    """插入 Word 目录域——打开文档后按 F9(或「更新目录」)即生成页码。"""
    p = doc.add_paragraph()
    r = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    inner = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "「右键此处 → 更新域」生成目录与页码"
    inner.append(t)
    fld.append(inner)
    r._element.addnext(fld)


def footer(doc, text):
    for sec in doc.sections:
        p = sec.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        _cn(r)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
def say(doc, title, lines):
    """对话示例框——本白皮书的主角:告诉读者「你可以这样说」。"""
    return box(doc, title, ["「" + ln + "」" for ln in lines], "blue")
def build(out_path: Path):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = CN_FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.4)
        sec.left_margin = sec.right_margin = Cm(2.8)
    for _ in range(5):
        doc.add_paragraph()
    para(doc, "心理学研究者的", size=20, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, font=CN_HEI, space_after=2)
    para(doc, "PsyClaw 使用白皮书", size=28, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, font=CN_HEI, space_after=14)
    para(doc, "以自然语言完成研究全流程", size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "—— 文献 · 统计分析 · 论文写作润色 ——", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=18)
    para(doc, "无需编程 · 无需记忆命令", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, space_after=40)
    para(doc, f"2026 年 7 月 · v{VER}", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
    page_break(doc)
    heading(doc, "关于本白皮书")
    para(doc, "本白皮书面向心理学及相关行为科学领域的研究者——研究生、博士后、"
              "青年教师与资深 PI——介绍如何用 PsyClaw 完成文献调研、统计分析"
              "与论文写作润色。")
    para(doc, "阅读本白皮书不需要编程基础,也不需要记住任何命令。安装完成后,"
              "全部交互均以自然语言进行:说明需求,由其执行;"
              "涉及写入文件、下载全文等操作时会先行征求同意。")
    para(doc, "我们假定你熟悉常见的心理学统计方法(t 检验、方差分析、回归、中介分析),"
              "日常可能使用 SPSS、JASP 或 R。分析脚本由 PsyClaw 生成并交给成熟统计库"
              "运行,你不必自己写。")
    heading(doc, "如何使用本白皮书", level=2)
    bullet(doc, "第 1 章:安装与配置。首次使用前请完整阅读。")
    bullet(doc, "第 2–4 章:文献 → 统计分析 → 写作润色,三个阶段各成一章,可按需查阅。")
    bullet(doc, "第 5 章:使用建议与能力边界。")
    bullet(doc, "附录:常用表述速查、量表配置、常见问题。")
    para(doc, "文中三类信息框:蓝色为对话示例(可直接照此表述),绿色为实用建议,"
              "黄色为风险提示。灰底为需在终端输入的命令——"
              "除安装配置外,正文中几乎不再涉及。", space_after=10)
    page_break(doc)
    heading(doc, "PsyClaw 能做什么")
    para(doc, "PsyClaw 承担研究流程中繁琐、易错而又必须准确的环节,"
              "使研究者得以专注于判断与思考。")
    table(doc, ["你想做的事", "你只需要说"],
          [["找某个主题近几年的文献", "帮我找工作倦怠与离职意向近三年的研究"],
           ["把找到的文章下载下来", "把这几篇下载下来"],
           ["确认参考文献没有编造", "帮我核一下这稿子里的引用是不是都真实存在"],
           ["做一个中介效应分析", "帮我做个中介分析,自变量是 X,因变量是 Y"],
           ["检查稿件是否符合规范", "帮我按 APA 规范检查一下这份稿子"],
           ["模拟审稿意见", "帮我审一下这篇,像 Nature 审稿人那样"],
           ["导出成 Word 投稿", "导出成心理学报格式的 Word"]], [5.6, 9.6])
    para(doc, "上表右列即为实际输入内容。无需参数与命令;"
              "表述不明确时,系统会主动追问。")
    box(doc, "三项基本保证",
        ["· 不编造文献。检索失败时如实告知,绝不以记忆内容拼凑书目。",
         "· 不编造数据。脚本未实际运行前,不提供含具体数值的「示例结果」。",
         "· 不擅自变更文件。写入、下载、修改文库前均先征求同意。"], "green")
    page_break(doc)
    heading(doc, "目录")
    for t in ["关于本白皮书", "PsyClaw 能做什么",
              "第一章　安装与配置",
              "第二章　文献检索与全文获取",
              "第三章　统计分析",
              "第四章　论文写作与润色",
              "第五章　使用建议与能力边界",
              "第六章　流程编排与能力扩展",
              "附录 A　常用表述速查",
              "附录 B　自定义量表配置",
              "附录 C　常见问题"]:
        p = para(doc, t, size=11, space_after=4)
        p.paragraph_format.left_indent = Cm(0.6)
    page_break(doc)
    return doc
def ch1(doc):
    heading(doc, "第一章　安装与配置")
    para(doc, "本章是全书唯一需要在终端输入命令的部分。配置完成后,"
              "后续交互均以自然语言进行。")
    heading(doc, "1.1　安装:国际网络环境", level=2)
    code(doc, ['uv tool install --python 3.12 \\',
               f'  "git+https://github.com/Exekiel179/psyclaw.git@v{VER}"'])
    para(doc, "若提示未找到 uv,请先安装:", space_after=2)
    code(doc, ["curl -LsSf https://astral.sh/uv/install.sh | sh"])
    heading(doc, "1.2　安装:国内网络环境", level=2)
    para(doc, "国内网络环境下直接执行上述命令通常会失败。建议使用一键脚本,"
              "该脚本会自动检测网络状况并切换至国内镜像:")
    code(doc, ["curl -fsSL https://exekiel179.github.io/psyclaw/install.sh | sh"])
    para(doc, "如需强制使用国内镜像,可在命令前加 PSYCLAW_CN=1。"
              "Windows 环境请以 PowerShell 执行同名的 install.ps1。")
    box(doc, "风险提示",
        ["一键脚本在 GitHub 不可达时会改用第三方镜像 gitclone.com。",
         "该镜像非官方渠道,代码完整性不作保证;能连 GitHub 时请优先用官方地址。"], "amber")
    heading(doc, "1.3　分发包:批量部署与无网环境", level=2)
    para(doc, "需为课题组批量部署,或目标机器无法联网时,使用分发包:")
    table(doc, ["文件", "适用场景"],
          [[f"psyclaw-{VER}-py3-none-any.whl", "常规离线安装(装时仍需拉一个依赖)"],
           [f"psyclaw-offline-{VER}.tar.gz", "完全无网:依赖全部打包,解压即装"]],
          [6.6, 8.6])
    para(doc, "全离线整包使用方式:拷贝至目标机器,解压后按平台执行其中的装机脚本——"
              "macOS 与 Linux 运行 install.sh(sh install.sh),Windows 则右键 install.ps1 "
              "选择「使用 PowerShell 运行」。两者仅需目标机器已具备 Python 3.11 及以上版本;"
              "整包为纯 Python 包,同一份文件在两类系统通用。")
    heading(doc, "1.4　首次配置", level=2)
    code(doc, ["psyclaw setup"])
    para(doc, "该命令引导完成功能板块选择、模型服务(provider)与密钥配置。"
              "配置一次即可,后续所有项目通用。")
    heading(doc, "1.5　四个入口", level=2)
    para(doc, "日常使用涉及以下四个入口。进入后即以自然语言交互。")
    table(doc, ["入口", "什么时候用"],
          [["psyclaw", "最常用。直接进入对话。"],
           ["psyclaw new 研究名", "新建研究。将建立独立目录,"
                                  "其目标与进度同其他研究相互隔离。"],
           ["psyclaw resume", "续接历史会话。不带参数时续接最近一次。"],
           ["psyclaw status", "查看当前研究进度:目标、待办事项、最近产出与后续建议。"]],
          [4.6, 10.6])
    box(doc, "实用建议",
        ["建议每项研究单独使用 psyclaw new 建立目录。",
         "系统将独立记录该研究的目标与进度,不会混入其他课题的上下文。"], "green")
    heading(doc, "1.6　使用须知", level=2)
    para(doc, "第一,表述不必精确。信息不足时系统会主动询问,而非擅自假设。")
    para(doc, "第二,涉及变更的操作均先征求同意。写入文件、下载全文、"
              "写入 Zotero 文库、打开浏览器等操作均会事先询问;"
              "同类操作经一次同意后,本次会话内可自动放行。")
    para(doc, "第三,产物存放位置固定:成稿位于 outputs/,图表位于 figures/,"
              "脚本位于 scripts/,笔记位于 notes/。原始数据目录 data/ 仅读取,不写入。")

    heading(doc, "1.7　权限与审批档位", level=2)
    para(doc, "PsyClaw 在两个维度上控制它能做什么:操作审批与文件访问。"
              "两者均可随时在对话中切换,切换后立即生效。")
    table(doc, ["档位", "切换方式", "含义"],
          [["审批 ask(默认)", "/approval ask",
            "命令执行、文件覆盖、工具副作用逐条征求同意"],
           ["审批 auto", "/approval auto 或 /yolo",
            "上述操作自动放行;仅命中红线的危险命令仍会询问"],
           ["访问 open(默认)", "/access open",
            "允许系统按需读取项目内文件"],
           ["访问 safe", "/access safe",
            "一切读取须由研究者以 @ 显式引用,系统不自行读取"]],
          [3.4, 3.6, 8.0])
    para(doc, "红线始终有效:删除类命令、强制推送、数据库删表等危险操作,"
              "无论处于何种档位都会要求确认。原始数据目录 data/ 在任何档位下均只读。")
    box(doc, "档位选择建议",
        ["初次使用或处理重要数据时,保持默认的 ask + open。",
         "批量重复性工作(如逐篇下载、批量导出)可临时切到 auto 提高效率,",
         "完成后切回。涉及未公开数据的分析,建议使用 safe 访问档位。"], "green")
    page_break(doc)
def ch2(doc):
    heading(doc, "第二章　文献检索与全文获取")
    para(doc, "文献是 AI 辅助研究中最易出现学术事故的环节。真正的风险不在于检索失败,"
              "而在于获得一份格式规范却并不存在的书目。本章所述流程,"
              "每一步均可核验。")
    heading(doc, "2.1　文献检索", level=2)
    say(doc, "对话示例",
        ["帮我找工作倦怠与离职意向近三年的研究,英文的,15 篇左右",
         "再找找中文文献,心理学报和心理科学上的",
         "这些太老了,只要 2024 年之后的"])
    para(doc, "系统将同时检索多个学术数据库(OpenAlex、Crossref、EuropePMC 等),"
              "去重后一并给出题录、年份、被引次数与 DOI。"
              "「近三年」以实际当前年份计算。")
    para(doc, "撰写综述时,沿引用网络检索更为有效:")
    say(doc, "对话示例",
        ["这篇很关键,帮我找找引用了它的后续研究",
         "顺着这篇的参考文献往回找源头"])
    box(doc, "实用建议",
        ["关键词检索易遗漏措辞不同的重要文献;以公认的经典文献为起点前后追溯,",
         "是更稳妥的综述路径。研究者只需指定种子文献。"], "green")
    heading(doc, "2.2　全文获取", level=2)
    say(doc, "对话示例",
        ["把上面这几篇下载下来", "第 3 篇和第 7 篇下载一下"])
    para(doc, "开放获取文献将直接下载至 outputs/pdfs 目录。"
              "遇付费墙时,系统不会就此中止,而是引导使用研究者本人的机构权限获取:")
    numbered(doc, "系统打开所在机构的访问入口(校园网代理或 LibKey);")
    numbered(doc, "在浏览器中以机构账号登录——该步骤须由研究者本人完成,"
                  "PsyClaw 不接触账号与密码;")
    numbered(doc, "点击页面上的下载按钮,保存至默认下载目录即可,"
                  "无需另存至指定位置,亦无需重命名;")
    numbered(doc, "返回对话说明已下载完成,系统将自动将文件收入项目,"
                  "并按「作者_年份_题名」重新命名。")
    para(doc, "若已安装配套浏览器扩展,第 3 步亦可省去,系统将直接取回文件。")
    box(doc, "关于付费墙文献",
        ["PsyClaw 不绕过任何付费墙,所使用的始终是研究者本人已有的机构权限。",
         "若所在单位未订阅该刊,系统将如实告知无法获取,不作其他尝试。"], "amber")
    para(doc, "使用 Zotero 者,可先在本人文库中检索:")
    say(doc, "对话示例",
        ["先在我的 Zotero 里看看有没有这几篇", "把这篇加到我的 Zotero 文库",
         "这篇我库里有全文,读一下它的方法部分"])
    heading(doc, "2.3　引用真实性核查", level=2)
    para(doc, "该步骤应在交稿前执行:")
    say(doc, "对话示例",
        ["帮我核一下这份稿子里的参考文献,是不是每条都真实存在"])
    para(doc, "系统将稿件中每条参考文献送至学术数据库逐条查证,"
              "并严格区分以下三种结果:")
    table(doc, ["结果", "含义"],
          [["查到了", "该文献真实存在,作者与年份都对得上"],
           ["查无此文", "疑似杜撰或著录有误,将明确标出待处理"],
           ["无法查证", "网络不通,或该类文献本就不被数据库收录"
                        "(如中文专著、内部报告)——只提示,不当作杜撰"]], [3.2, 11.4])
    box(doc, "该步骤的必要性",
        ["格式规整、卷期页码齐全、正文引用与文末列表完全一致的虚构条目,",
         "仅凭人工查阅与格式检查无法识别,须经数据库实际查证。",
         "「查无此文」与「无法查证」为两类不同结果,后者不会误判未被收录的中文文献。"], "blue")
    page_break(doc)
def ch3(doc):
    heading(doc, "第三章　统计分析")
    para(doc, "研究者无需编写代码。说明研究问题、变量与数据位置后,"
              "系统将生成分析脚本、交由成熟统计库运行,并解读结果。"
              "脚本保存于 scripts/ 目录,可供他人复现验证。")
    para(doc, "首次进行统计分析前,需安装统计库(仅需一次):", space_after=2)
    code(doc, ['pip install "psyclaw[stats]"'])
    heading(doc, "3.1　分析前的准备", level=2)
    para(doc, "分析结果的可信度,很大程度上取决于分析计划确定于接触数据之前还是之后。"
              "因此第一步并非运行数据,而是记录分析计划。")
    say(doc, "对话示例",
        ["我要研究工作倦怠在工作满意度和离职意向之间的中介作用,帮我把研究准备清单过一遍",
         "这个假设是确证性的,帮我先声明下来",
         "我打算用中介分析,样本量需要多少?"])
    para(doc, "系统将逐项询问:结局如何定义、自变量如何操作化、排除标准为何、"
              "拟采用何种检验。完成后该清单将保存下来,"
              "供撰写方法部分、进行预注册与稿件检查时使用。")
    box(doc, "实用建议",
        ["「先声明,后分析」并非形式要求。已声明的假设属确证性,",
         "未经声明所得结果属探索性——二者在论文中的表述方式与证据分量截然不同。",
         "PsyClaw 将守住这一界线,避免探索性结果被表述为确证性发现。"], "green")
    heading(doc, "3.2　数据清理与量表计分", level=2)
    say(doc, "对话示例",
        ["数据在 data/clean/survey.csv,帮我看看这份数据什么情况",
         "帮我按 MBI 的计分规则算总分,第 13、14 题是反向题",
         "有没有明显乱答的被试?比如一直选同一个选项的"])
    para(doc, "系统可执行反向题翻转、子量表求和、缺失值处理,并识别可疑作答模式。"
              "量表计分规则由研究者提供一次(见附录 B),此后持续可用。")
    heading(doc, "3.3　执行分析", level=2)
    say(doc, "对话示例",
        ["帮我做个中介分析,自变量工作满意度,中介变量倦怠,因变量离职意向",
         "两组比较一下,看看性别有没有差异",
         "帮我看看这些数据适不适合做回归,前提假设满足吗",
         "把这个分析写成脚本,我要能自己重跑"])
    para(doc, "系统将生成脚本并运行,随后以日常语言解读结果:"
              "先说明结果的实际含义,再给出统计数值。")
    box(doc, "统计报告规范",
        ["· 必报效应量与 95% 置信区间,不止于 p 值;",
         "· p 值大于 .05 即如实报告不显著,不使用「边缘显著」「显著趋势」等表述;",
         "· 相关不等于因果,涉及因果的表述附带识别假设;",
         "· 脚本未实际运行前,不提供含具体数值的「示例结果」。"], "blue")
    heading(doc, "3.4　完整流程", level=2)
    para(doc, "如需系统串联执行完整分析流程:")
    say(doc, "对话示例",
        ["帮我把这份数据的分析从头走一遍:先描述,再主分析,最后做稳健性检验",
         "我要做元分析,效应量表在这个文件里",
         "这是访谈转录稿,帮我做主题分析"])
    page_break(doc)
def ch4(doc):
    heading(doc, "第四章　论文写作与润色")
    heading(doc, "4.1　撰写初稿", level=2)
    say(doc, "对话示例",
        ["根据上面的分析结果,帮我写方法部分",
         "结果部分写一下,按 APA 格式报告",
         "帮我把讨论部分的第二段改得紧凑一些",
         "这段话太啰嗦了,精简到 150 字以内"])
    para(doc, "撰写过程遵循两条规则:统计数值仅使用实际运行所得结果,"
              "缺失处以「(待回填)」占位并在文末列出清单;"
              "引用仅使用检索所得真实文献,缺失处标注「(待补引)」,"
              "检索到后再行填补,不以看似合理的条目填充。")
    heading(doc, "4.2　投稿前检查", level=2)
    say(doc, "对话示例",
        ["帮我按 APA 规范检查一下这份稿子",
         "对照 JARS 清单看看还缺什么",
         "帮我审一下这篇,像 Nature 审稿人那样,尽量挑刺"])
    para(doc, "检查将一次性汇总以下结果:JARS 清单是否存在遗漏、"
              "参考文献是否均真实存在、格式是否符合目标期刊要求、"
              "结论是否具备相应的分析支撑。")
    box(doc, "内置审稿流程",
        ["审稿模拟基于从 1287 篇 Nature 审稿报告中提炼的流程:",
         "分 12 个维度逐一审查,并给出可执行的修改建议。",
         "投稿前先行自查,可显著降低因基础问题被拒的风险。"], "green")
    heading(doc, "4.3　图表", level=2)
    say(doc, "对话示例",
        ["帮我画一张中介模型的路径图",
         "把这三组的均值画成柱状图,APA 风格",
         "这张图的纵轴是不是截断了?看着有点夸大"])
    para(doc, "图表保存至 figures/ 目录,中文标签正常显示。"
              "系统亦会提示截断坐标轴等易造成误导的呈现方式。")
    heading(doc, "4.4　导出与投稿", level=2)
    say(doc, "对话示例",
        ["导出成 Word,APA7 格式",
         "按心理学报的格式导出",
         "这篇准备投心理科学,帮我看看格式要求对不对得上"])
    para(doc, "导出的 Word 文件版式确定:字体、行距、三级标题、"
              "参考文献悬挂缩进与页码位置均按目标期刊要求排布,图片真实嵌入。"
              "当前支持 APA7、心理学报、心理科学三种格式。")
    page_break(doc)
def ch5(doc):
    heading(doc, "第五章　使用建议与能力边界")
    heading(doc, "5.1　使用建议", level=2)
    bullet(doc, "每项课题使用 psyclaw new 单独建立目录,进度与目标相互隔离。")
    bullet(doc, "开始时先说明研究目标,系统将予以记录,"
                "后续建议均围绕该目标给出。")
    bullet(doc, "过程中可询问当前进展,系统将汇报进度并给出后续建议。")
    bullet(doc, "对系统所做变更存疑时,可直接询问其改动了哪些文件。")
    bullet(doc, "若系统声称无法完成某项操作,而该功能理应存在,"
                "通常是未匹配到对应能力;更换表述重新提出即可。")
    heading(doc, "5.2　能力边界", level=2)
    para(doc, "以下能力边界请在使用前知悉:")
    numbered(doc, "不代为构思研究问题。系统可协助厘清表述、指出逻辑漏洞,"
                  "但选题仍由研究者决定。")
    numbered(doc, "不内置量表库。常用量表的条目与计分规则需由研究者提供一次"
                  "(见附录 B),此后持续可用。")
    numbered(doc, "无法获取所在单位未订阅的文献。系统使用的是研究者本人的机构权限,"
                  "权限不足时不作其他尝试。")
    numbered(doc, "不承担学术责任。全部统计解读与文献引用最终均须研究者本人复核"
                  "——论文的责任作者是研究者。")
    heading(doc, "5.3　风险提示", level=2)
    box(doc, "务必留意",
        ["· 关键引用须阅读原文。查证仅能确认文献真实存在,",
         "  不能确认其支持所提出的论点;",
         "· 涉及被试隐私的数据,发送至模型前须确认符合所在伦理委员会要求;",
         "· 原始数据须自行备份。PsyClaw 不写入 data/ 目录,但备份责任在研究者;",
         "· AI 生成的全部内容,发表后署名的是研究者本人。"], "amber")
    heading(doc, "5.4　问题排查", level=2)
    para(doc, "可在对话中直接说明问题,例如某步骤报错或结果异常,"
              "系统通常可自行定位。若属环境问题,在终端执行:", space_after=2)
    code(doc, ["psyclaw doctor"])
    para(doc, "仍无法解决时,请将当时的原始输出提交至项目页面:"
              "https://github.com/Exekiel179/psyclaw/issues")
    page_break(doc)
def ch6(doc):
    heading(doc, "第六章　流程编排与能力扩展")
    para(doc, "前几章介绍的是单步操作。当研究进入需要多步骤串联的阶段时,"
              "可使用流程编排;当内置能力不足时,可通过 skill 与 MCP 扩展。")
    heading(doc, "6.1　四类研究流程", level=2)
    para(doc, "流程会自动串联「设计 → 执行 → 产出 → 评审」各环节,"
              "并在每个环节留下可复现的记录。可在对话中直接说明,"
              "亦可在终端指定流程类型。")
    table(doc, ["流程", "适用研究类型", "主要环节"],
          [["analysis", "实证数据分析",
            "数据画像 → 设计核对 → 推荐分析并生成脚本 → 结果评审"],
           ["meta", "元分析",
            "效应量表校验 → 生成可复现脚本 → 撰写 → 评审"],
           ["literature", "文献综述",
            "研究准备 → 检索 → PRISMA 筛选 → 综述合成 → 评审"],
           ["qualitative", "质性研究",
            "载入转录稿 → 设计 → 主题分析 → COREQ 报告 → 评审"]],
          [2.6, 3.4, 9.0])
    say(doc, "对话示例",
        ["帮我把这份数据的实证分析流程完整走一遍,数据在 data/clean/survey.csv",
         "我要做一个关于正念干预的元分析,效应量表已经整理好了",
         "帮我做工作倦怠这个主题的文献综述,要按 PRISMA 流程"])
    para(doc, "对应的终端形式(需要精确控制时使用):", space_after=2)
    code(doc, ["psyclaw run analysis data/clean/survey.csv",
               "psyclaw run literature --topic 工作倦怠与离职意向",
               "psyclaw run qualitative data/transcripts/"])
    para(doc, "三个常用选项:--confirm-each 在每步完成后暂停确认;"
              "--resume 从上次中断处继续;--exploratory 允许跳过未完成的前置检查,"
              "产出将被明确标注为探索性。")
    box(doc, "关于自动推进",
        ["除上述四类流程外,还可让系统根据项目当前状态自主决定下一步:",
         "在对话中说明「按现在的进度继续推进」即可。",
         "强制性检查与不可逆决策仍会暂停并征求同意,不会一路跑到底。"], "blue")
    heading(doc, "6.2　skill:把方法学流程装进来", level=2)
    para(doc, "skill 是一份结构化的方法学流程说明。系统内置若干 skill,"
              "在对话涉及相应主题时自动调用,无需手动指定。")
    table(doc, ["内置 skill", "用途"],
          [["nature-review", "Nature 级同行评审与回复信撰写"],
           ["sample-size", "样本量估算(功效分析)"],
           ["confound-control", "无关变量与混淆变量的控制流程"],
           ["analysis-planning", "实证分析计划制定"],
           ["lit-review", "文献综述与证据合成"],
           ["pingouin", "统计函数选择指南"],
           ["paper-review-gates", "稿件评审与质量把关"]], [4.4, 10.8])
    para(doc, "亦可安装第三方 skill。在终端执行:", space_after=2)
    code(doc, ["psyclaw skill install https://github.com/用户名/仓库名"])
    para(doc, "安装后的 skill 位于项目的 .claude/skills 目录,系统会自动发现并纳入调用范围。"
              "查看当前已装载的 skill,在对话中询问即可。")
    box(doc, "安全提示",
        ["skill 安装仅接受 https 协议的 GitHub 地址。",
         "第三方 skill 的内容会进入系统提示并影响模型行为,",
         "请确认来源可信后再安装。"], "amber")
    heading(doc, "6.3　MCP:接入外部统计软件与服务", level=2)
    para(doc, "MCP(Model Context Protocol)用于把外部工具接入对话。"
              "配置完成后,这些工具与内置能力一样可被直接调用。")
    table(doc, ["MCP 服务", "用途"],
          [["pystat", "Python 统计计算(psyclaw 的统计委托通道)"],
           ["r-mcp", "调用本机 R 环境"],
           ["spss-mcp / stata-mcp / mplus-mcp", "调用 SPSS、Stata、Mplus"],
           ["mne-mcp", "脑电与脑磁数据处理"],
           ["zotero-mcp", "Zotero 文献管理"],
           ["osf-mcp", "OSF 预注册与资料托管"],
           ["sequential-thinking", "复杂问题的分步推理"]], [5.2, 10.0])
    para(doc, "查看目录与启用状态:", space_after=2)
    code(doc, ["psyclaw mcp        # 查看 MCP 目录与当前启用状态",
               "psyclaw config     # 配置向导中可完成密钥与路径设置"])
    para(doc, "已有 SPSS、Mplus 使用习惯的研究者,可通过对应 MCP 继续使用原有软件,"
              "由 PsyClaw 负责流程串联与结果归档。")
    box(doc, "实用建议",
        ["MCP 采用惰性加载:未实际调用时不会占用上下文,",
         "因此可以按需配置多个而不必担心影响日常对话的响应速度。"], "green")
    page_break(doc)

def appendices(doc):
    heading(doc, "附录 A　常用表述速查")
    para(doc, "下表左列为常见需求,右列为可直接采用的表述。")
    table(doc, ["想做的事", "可以这样说"],
          [["开始一个新研究", "我要研究 X 对 Y 的影响,先把研究准备清单过一遍"],
           ["找文献", "帮我找 X 近三年的研究,15 篇左右"],
           ["顺着引用找", "帮我找找引用了这篇的后续研究"],
           ["下载全文", "把上面这几篇下载下来"],
           ["用机构权限拿付费文章", "这篇有付费墙,帮我走机构权限"],
           ["查我的 Zotero", "先在我的 Zotero 里看看有没有"],
           ["核查引用真伪", "核一下这稿子里的参考文献是不是都真实存在"],
           ["看数据情况", "数据在这个文件,帮我看看什么情况"],
           ["量表计分", "按这个量表的规则算总分,第 N 题是反向题"],
           ["做分析", "帮我做个中介分析,自变量 X,中介 M,因变量 Y"],
           ["检查前提假设", "这些数据适合做回归吗,前提假设满足吗"],
           ["写稿", "根据分析结果帮我写方法部分"],
           ["润色", "这段太啰嗦,精简到 150 字"],
           ["规范检查", "按 APA 规范检查一下这份稿子"],
           ["模拟审稿", "帮我审一下这篇,尽量挑刺"],
           ["画图", "把这三组均值画成柱状图,APA 风格"],
           ["导出", "导出成心理学报格式的 Word"],
           ["看进度", "现在进展到哪了"],
           ["走完整分析流程", "帮我把这份数据的实证分析流程完整走一遍"],
           ["做文献综述流程", "帮我做这个主题的文献综述,按 PRISMA 流程"],
           ["自主推进", "按现在的进度继续推进"],
           ["放宽审批", "/approval auto(批量操作时用,完成后切回)"],
           ["收紧文件访问", "/access safe(处理未公开数据时用)"]], [4.4, 10.8])
    heading(doc, "附录 B　自定义量表配置")
    para(doc, "PsyClaw 不内置量表库——常用量表版本众多、计分规则各异,"
              "内置不完整的量表库反而易致误用。研究者提供一次所用量表,此后持续可用。")
    para(doc, "最简便的方式是在对话中直接说明:", space_after=2)
    say(doc, "对话示例",
        ["我要用 MBI 通用版,16 题,0–6 计分,第 13、14 题反向,"
         "分三个维度:情绪衰竭 1–5 题,玩世不恭 6–9 题,职业效能 10–16 题。"
         "帮我把它加到量表定义里"])
    para(doc, "系统将生成量表定义文件保存至项目中。亦可自行在"
              " .psyclaw/scales/ 目录下编写,格式如下:", space_after=2)
    code(doc, ["- id: mbi-gs",
               "  name: Maslach 倦怠量表(通用版)",
               "  items: 16",
               "  scale_min: 0",
               "  scale_max: 6",
               "  reverse: [13, 14]",
               "  subscales:",
               "    情绪衰竭: [1, 2, 3, 4, 5]",
               "    玩世不恭: [6, 7, 8, 9]",
               "    职业效能: [10, 11, 12, 13, 14, 15, 16]"])
    box(doc, "实用建议",
        ["量表定义编写一次即可跨研究复用。建议将课题组常用量表一次性配置完成,",
         "并在 name 字段注明所用中文修订版本及来源文献——",
         "撰写方法部分时需引用。"], "green")
    heading(doc, "附录 C　常见问题")
    qa = [
        ("一定要联网吗?",
         "文献检索与引用核查需联网;撰写、检查、导出、计分均可离线进行。"
         "调用模型本身需要网络连接。"),
        ("我的数据会被上传吗?",
         "PsyClaw 不会主动读取原始数据目录。在对话中明确指定读取的文件,"
         "其内容将随请求发送至所配置的模型服务——涉及被试隐私时请审慎处理,"
         "并确认符合伦理委员会要求。"),
        ("能替代 SPSS 或 JASP 吗?",
         "定位不同。其长处在于串联完整流程并留下可复现记录;"
         "点选式软件的即时探索体验则非其所长。二者可并行使用。"),
        ("生成的稿子能直接投吗?",
         "不能。全部统计解读与文献引用均须研究者本人复核。"
         "其价值在于将易错环节转化为可核验的流程,而非代为承担责任。"),
        ("它会不会编造文献?",
         "不会。检索失败时如实告知,并给出后续可行方案。"
         "交稿前可要求其将每条参考文献送数据库逐条核验。"),
        ("换电脑了怎么办?",
         "重新安装即可。配置与量表定义随项目目录保存,"
         "拷贝项目目录至新机器即可继续使用。"),
    ]
    for q, a in qa:
        para(doc, "问:" + q, bold=True, space_after=2)
        para(doc, "答:" + a, space_after=8)
    para(doc, f"PsyClaw v{VER} · 源码与问题反馈:"
              "https://github.com/Exekiel179/psyclaw · MIT 许可",
         size=9.5, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
def main():
    out = ROOT / "docs" / f"PsyClaw使用白皮书_v{VER}.docx"
    doc = build(out)
    ch1(doc)
    ch2(doc)
    ch3(doc)
    ch4(doc)
    ch5(doc)
    ch6(doc)
    appendices(doc)
    footer(doc, f"PsyClaw 使用白皮书 · v{VER}")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"✓ {out}")
    return out
if __name__ == "__main__":
    main()
